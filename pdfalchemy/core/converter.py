"""
Format converters: PDF → Editable formats with layout preservation.
"""
from typing import List, Optional
from pathlib import Path
import json
import logging

from .extractor import StreamingExtractor, PageLayout, TextBlock

logger = logging.getLogger(__name__)


class BaseConverter:
    def __init__(self, extractor: StreamingExtractor):
        self.extractor = extractor

    def convert(self, output_path: str, **options) -> str:
        raise NotImplementedError


class DOCXConverter(BaseConverter):
    """Convert PDF to editable Word document with style mapping."""

    def convert(self, output_path: str, **options) -> str:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        for page_layout in self.extractor.stream_pages():
            for block in page_layout.blocks:
                if block.block_type == "header":
                    continue
                elif block.block_type == "footer":
                    continue
                elif block.block_type == "heading":
                    doc.add_heading(block.text, level=1)
                elif block.block_type == "caption":
                    doc.add_paragraph(block.text, style='Caption')
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(block.text)
                    run.font.size = Pt(block.size)
                    run.font.name = block.font.split("+")[-1]

            if options.get("page_breaks", True):
                doc.add_page_break()

        doc.save(output_path)
        return output_path


class HTMLConverter(BaseConverter):
    """Convert PDF to semantic HTML with CSS styling."""

    def convert(self, output_path: str, **options) -> str:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(
            "<!DOCTYPE html><html><head></head><body></body></html>", 
            "html.parser"
        )
        head = soup.head
        body = soup.body

        style = soup.new_tag("style")
        style.string = """
            .pdf-page { margin: 20px auto; max-width: 800px; border: 1px solid #ccc; padding: 40px; }
            .header { font-size: 0.8em; color: #666; border-bottom: 1px solid #eee; margin-bottom: 20px; }
            .footer { font-size: 0.8em; color: #666; border-top: 1px solid #eee; margin-top: 20px; }
            .heading { font-size: 1.5em; font-weight: bold; margin: 20px 0 10px; }
            .body { margin: 10px 0; line-height: 1.6; }
            .caption { font-size: 0.9em; color: #555; font-style: italic; }
            table { border-collapse: collapse; width: 100%; margin: 15px 0; }
            td, th { border: 1px solid #ddd; padding: 8px; }
        """
        head.append(style)

        for page_layout in self.extractor.stream_pages():
            page_div = soup.new_tag("div", **{"class": "pdf-page"})
            page_div["data-page-num"] = str(page_layout.page_num)

            for block in page_layout.blocks:
                if block.block_type == "header":
                    tag = soup.new_tag("div", **{"class": "header"})
                elif block.block_type == "footer":
                    tag = soup.new_tag("div", **{"class": "footer"})
                elif block.block_type == "heading":
                    tag = soup.new_tag("h2")
                elif block.block_type == "caption":
                    tag = soup.new_tag("p", **{"class": "caption"})
                else:
                    tag = soup.new_tag("p", **{"class": "body"})

                tag.string = block.text
                page_div.append(tag)

            for table in page_layout.tables:
                table_tag = soup.new_tag("table")
                for row_data in table["cells"]:
                    tr = soup.new_tag("tr")
                    for cell_text in row_data:
                        td = soup.new_tag("td")
                        td.string = cell_text
                        tr.append(td)
                    table_tag.append(tr)
                page_div.append(table_tag)

            body.append(page_div)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(str(soup.prettify()))

        return output_path


class MarkdownConverter(BaseConverter):
    """Convert PDF to Markdown with structure preservation."""

    def convert(self, output_path: str, **options) -> str:
        lines = []

        for page_layout in self.extractor.stream_pages():
            for block in page_layout.blocks:
                if block.block_type in ("header", "footer"):
                    continue
                elif block.block_type == "heading":
                    lines.append(f"\n## {block.text}\n")
                elif block.block_type == "caption":
                    lines.append(f"*{block.text}*\n")
                else:
                    lines.append(f"{block.text}\n")

            for table in page_layout.tables:
                lines.append("\n")
                for i, row in enumerate(table["cells"]):
                    lines.append("| " + " | ".join(row) + " |")
                    if i == 0:
                        lines.append("| " + " | ".join(["---"] * len(row)) + " |")
                lines.append("\n")

            lines.append("\n---\n")

        content = "\n".join(lines)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)

        return output_path


class JSONConverter(BaseConverter):
    """Convert PDF to structured JSON for data processing."""

    def convert(self, output_path: str, **options) -> str:
        pages = []

        for page_layout in self.extractor.stream_pages():
            page_dict = {
                "page_num": page_layout.page_num,
                "dimensions": {
                    "width": page_layout.width, 
                    "height": page_layout.height
                },
                "blocks": [
                    {
                        "text": b.text,
                        "bbox": b.bbox,
                        "type": b.block_type,
                        "font": b.font,
                        "size": b.size
                    }
                    for b in page_layout.blocks
                ],
                "tables": page_layout.tables,
                "image_count": len(page_layout.images)
            }
            pages.append(page_dict)

        result = {
            "document": {
                "total_pages": len(pages),
                "pages": pages
            }
        }

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result, f, indent=2, ensure_ascii=False)

        return output_path


class ConverterFactory:
    _converters = {
        "docx": DOCXConverter,
        "html": HTMLConverter,
        "md": MarkdownConverter,
        "markdown": MarkdownConverter,
        "json": JSONConverter,
    }

    @classmethod
    def get_converter(cls, format_type: str, extractor: StreamingExtractor):
        format_type = format_type.lower().lstrip(".")
        if format_type not in cls._converters:
            raise ValueError(
                f"Unsupported format: {format_type}. "
                f"Supported: {list(cls._converters.keys())}"
            )
        return cls._converters[format_type](extractor)
